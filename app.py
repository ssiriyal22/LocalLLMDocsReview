#!/usr/bin/env python3
"""
Local LLM Document Review Application
A self-contained application for reviewing documents and PDFs using open source LLMs
"""

import os
import sys
import json
import asyncio
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
import tempfile
import shutil
import subprocess

# Third-party imports (will be installed via requirements)
try:
    import streamlit as st
    import PyPDF2
    import docx
    import requests
    from sentence_transformers import SentenceTransformer
    import faiss
    import numpy as np
    import ollama
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain_community.vectorstores import FAISS
    from langchain_community.embeddings import HuggingFaceEmbeddings
    from langchain.schema import Document
except ImportError as e:
    print(f"Missing required packages. Please run the installer first: {e}")
    sys.exit(1)

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles document loading and processing"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.txt', '.docx', '.md']
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF files"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            return ""
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error reading text file: {e}")
            return ""
    
    def process_document(self, file_path: str) -> str:
        """Process document based on file extension"""
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_ext == '.docx':
            return self.extract_text_from_docx(file_path)
        elif file_ext in ['.txt', '.md']:
            return self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

class VectorStore:
    """Handles document vectorization and similarity search"""
    
    def __init__(self, model_name: str = "all-MiniLM-L6-v2"):
        self.embeddings = HuggingFaceEmbeddings(model_name=model_name)
        self.vector_store = None
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len
        )
    
    def create_vector_store(self, documents: List[str]) -> None:
        """Create vector store from documents"""
        # Split documents into chunks
        all_chunks = []
        for doc_text in documents:
            chunks = self.text_splitter.split_text(doc_text)
            all_chunks.extend([Document(page_content=chunk) for chunk in chunks])
        
        # Create vector store
        if all_chunks:
            self.vector_store = FAISS.from_documents(all_chunks, self.embeddings)
            logger.info(f"Created vector store with {len(all_chunks)} chunks")
    
    def similarity_search(self, query: str, k: int = 4) -> List[str]:
        """Search for similar documents"""
        if not self.vector_store:
            return []
        
        docs = self.vector_store.similarity_search(query, k=k)
        return [doc.page_content for doc in docs]

class LLMManager:
    """Manages local LLM interactions via Ollama"""
    
    def __init__(self):
        self.available_models = []
        self.current_model = None
        self.check_ollama_status()
    
    def check_ollama_status(self) -> bool:
        """Check if Ollama is running"""
        try:
            response = requests.get("http://localhost:11434/api/tags", timeout=5)
            if response.status_code == 200:
                models_data = response.json()
                self.available_models = [model['name'] for model in models_data.get('models', [])]
                logger.info(f"Found {len(self.available_models)} models: {self.available_models}")
                return True
        except requests.exceptions.ConnectionError:
            logger.warning("Ollama service not running on localhost:11434")
        except requests.exceptions.Timeout:
            logger.warning("Ollama service timeout")
        except Exception as e:
            logger.warning(f"Ollama not accessible: {e}")
        return False
    
    def pull_model(self, model_name: str) -> bool:
        """Pull a model using Ollama"""
        try:
            subprocess.run(['ollama', 'pull', model_name], check=True)
            self.check_ollama_status()  # Refresh available models
            return True
        except subprocess.CalledProcessError as e:
            logger.error(f"Failed to pull model {model_name}: {e}")
            return False
    
    def generate_response(self, prompt: str, context: List[str] = None) -> str:
        """Generate response using the current model"""
        if not self.current_model:
            return "No model selected. Please select a model first."
        
        # Prepare the full prompt with context
        full_prompt = prompt
        if context:
            context_text = "\n".join(context)
            full_prompt = f"Context:\n{context_text}\n\nQuestion: {prompt}\n\nAnswer:"
        
        try:
            response = ollama.generate(
                model=self.current_model,
                prompt=full_prompt,
                stream=False
            )
            return response['response']
        except Exception as e:
            logger.error(f"Error generating response: {e}")
            return f"Error: {str(e)}"

class DocumentReviewApp:
    """Main application class"""
    
    def __init__(self):
        self.doc_processor = DocumentProcessor()
        self.vector_store = VectorStore()
        self.llm_manager = LLMManager()
        self.documents = {}
        self.setup_page_config()
        
        # Initialize session state for model selection
        if 'current_model' not in st.session_state:
            st.session_state.current_model = None
        
        # Sync session state with LLM manager
        if st.session_state.current_model:
            self.llm_manager.current_model = st.session_state.current_model
    
    def setup_page_config(self):
        """Configure Streamlit page"""
        st.set_page_config(
            page_title="Local LLM Document Review",
            page_icon="📄",
            layout="wide",
            initial_sidebar_state="expanded"
        )
    
    def sidebar_model_management(self):
        """Handle model management in sidebar"""
        st.sidebar.header("🤖 Model Management")
        
        # Check Ollama status with retry button
        col1, col2 = st.sidebar.columns([3, 1])
        with col1:
            if st.button("🔄 Refresh", help="Check Ollama status"):
                st.rerun()
        
        ollama_running = self.llm_manager.check_ollama_status()
        
        if not ollama_running:
            st.sidebar.error("❌ Ollama not running")
            st.sidebar.markdown("**To fix this:**")
            st.sidebar.code("ollama serve", language="bash")
            st.sidebar.markdown("**Or try:**")
            if st.sidebar.button("🚀 Try to Start Ollama"):
                try:
                    import subprocess
                    subprocess.Popen(['ollama', 'serve'], 
                                   stdout=subprocess.DEVNULL, 
                                   stderr=subprocess.DEVNULL)
                    st.sidebar.info("Attempting to start Ollama... Please wait 5 seconds and refresh.")
                except Exception as e:
                    st.sidebar.error(f"Failed to start Ollama: {e}")
            return
        
        st.sidebar.success("✅ Ollama is running")
        
        # Show current model status with more prominence
        if self.llm_manager.current_model:
            st.sidebar.success(f"🎯 **Active Model:** {self.llm_manager.current_model}")
        else:
            st.sidebar.error("❌ **No Model Selected**")
            
        # Auto-select the only available model if there's just one
        if (len(self.llm_manager.available_models) == 1 and 
            not self.llm_manager.current_model):
            self.llm_manager.current_model = self.llm_manager.available_models[0]
            st.session_state.current_model = self.llm_manager.current_model
            st.sidebar.success(f"🎯 Auto-selected: {self.llm_manager.current_model}")
            st.rerun()
        
        # Available models section
        if self.llm_manager.available_models:
            st.sidebar.subheader("📋 Available Models")
            
            # Show available models as buttons for easier selection
            st.sidebar.markdown("**Click to select a model:**")
            for model in self.llm_manager.available_models:
                col1, col2 = st.sidebar.columns([3, 1])
                with col1:
                    is_current = model == self.llm_manager.current_model
                    button_text = f"🎯 {model}" if is_current else f"⚪ {model}"
                    if st.button(button_text, key=f"select_{model}", 
                               use_container_width=True,
                               type="primary" if is_current else "secondary"):
                        if not is_current:
                            self.llm_manager.current_model = model
                            st.sidebar.success(f"✅ Selected: {model}")
                            # Force session state update
                            if 'current_model' not in st.session_state:
                                st.session_state.current_model = model
                            else:
                                st.session_state.current_model = model
                            st.rerun()
                with col2:
                    if is_current:
                        st.sidebar.markdown("**✓**")
            
            # Also provide traditional dropdown as backup
            st.sidebar.markdown("---")
            st.sidebar.markdown("**Or use dropdown:**")
            
            # Initialize session state for model selection
            if 'selected_model_dropdown' not in st.session_state:
                st.session_state.selected_model_dropdown = self.llm_manager.current_model or ""
            
            selected_model = st.sidebar.selectbox(
                "Model:",
                ["Select a model..."] + self.llm_manager.available_models,
                index=0 if not self.llm_manager.current_model else 
                      (self.llm_manager.available_models.index(self.llm_manager.current_model) + 1),
                key="model_dropdown"
            )
            
            if selected_model != "Select a model..." and selected_model != self.llm_manager.current_model:
                self.llm_manager.current_model = selected_model
                st.session_state.current_model = selected_model
                st.sidebar.success(f"✅ Selected: {selected_model}")
                st.rerun()
                
        else:
            st.sidebar.warning("⚠️ No models installed")
            st.sidebar.info("Install a model using the section below.")
        
        # Pull new models
        st.sidebar.subheader("⬇️ Install New Model")
        recommended_models = [
            "llama3.2:1b",  # Fastest, smallest
            "llama3.2:3b",  # Recommended balance
            "phi3:mini",    # Microsoft efficient
            "mistral:7b",   # High quality
            "codellama:7b"  # For technical docs
        ]
        
        model_descriptions = {
            "llama3.2:1b": "⚡ Fast & lightweight (1.3GB)",
            "llama3.2:3b": "🎯 Recommended balance (2.0GB)", 
            "phi3:mini": "🔧 Microsoft efficient (2.2GB)",
            "mistral:7b": "🌟 High quality (4.1GB)",
            "codellama:7b": "💻 For code/technical (3.8GB)"
        }
        
        new_model = st.sidebar.selectbox(
            "Choose Model:",
            [""] + recommended_models,
            format_func=lambda x: model_descriptions.get(x, x) if x else "Select a model...",
            key="new_model_select"
        )
        
        custom_model = st.sidebar.text_input("Or enter custom model name:", 
                                           placeholder="e.g., gemma:2b")
        
        model_to_pull = custom_model if custom_model else new_model
        
        if st.sidebar.button("📥 Install Model", disabled=not model_to_pull):
            with st.sidebar.spinner(f"Installing {model_to_pull}... This may take several minutes."):
                if self.llm_manager.pull_model(model_to_pull):
                    st.sidebar.success(f"✅ {model_to_pull} installed!")
                    # Auto-select the newly installed model
                    self.llm_manager.current_model = model_to_pull
                    st.rerun()
                else:
                    st.sidebar.error(f"❌ Failed to install {model_to_pull}")
        
        # Model management options
        if self.llm_manager.available_models:
            st.sidebar.subheader("🗂️ Manage Models")
            if st.sidebar.button("🔄 Refresh Model List"):
                self.llm_manager.check_ollama_status()
                st.rerun()
            
            # Show model sizes/info
            with st.sidebar.expander("📊 Model Information"):
                for model in self.llm_manager.available_models:
                    st.sidebar.text(f"• {model}")
                st.sidebar.markdown("**Remove models:**")
                st.sidebar.code("ollama rm model_name", language="bash")
    
    def document_upload_section(self):
        """Handle document upload"""
        st.header("📄 Document Upload")
        
        uploaded_files = st.file_uploader(
            "Upload documents (PDF, DOCX, TXT, MD)",
            type=['pdf', 'docx', 'txt', 'md'],
            accept_multiple_files=True,
            key="doc_uploader"
        )
        
        if uploaded_files:
            for uploaded_file in uploaded_files:
                if uploaded_file.name not in self.documents:
                    # Save uploaded file temporarily
                    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp_file:
                        tmp_file.write(uploaded_file.getvalue())
                        tmp_path = tmp_file.name
                    
                    try:
                        # Process document
                        text = self.doc_processor.process_document(tmp_path)
                        if text.strip():
                            self.documents[uploaded_file.name] = text
                            st.success(f"✅ Processed: {uploaded_file.name}")
                        else:
                            st.error(f"❌ No text extracted from: {uploaded_file.name}")
                    except Exception as e:
                        st.error(f"❌ Error processing {uploaded_file.name}: {str(e)}")
                    finally:
                        # Clean up temp file
                        os.unlink(tmp_path)
            
            # Create vector store from all documents
            if self.documents:
                with st.spinner("Creating vector index..."):
                    self.vector_store.create_vector_store(list(self.documents.values()))
                st.success(f"✅ Vector index created for {len(self.documents)} documents")
    
    def chat_interface(self):
        """Main chat interface"""
        st.header("💬 Document Review Chat")
        
        # Check prerequisites
        if not self.llm_manager.check_ollama_status():
            st.error("❌ Ollama is not running. Please start Ollama first:")
            st.code("ollama serve", language="bash")
            return
        
        if not self.llm_manager.current_model:
            st.warning("⚠️ Please select a model from the sidebar first.")
            if self.llm_manager.available_models:
                st.info("💡 Available models found in sidebar. Please select one to continue.")
            else:
                st.info("💡 No models installed. Please install a model from the sidebar first.")
            return
        
        if not self.documents:
            st.warning("⚠️ Please upload documents first.")
            st.info("💡 Use the document upload section to add PDF, DOCX, TXT, or MD files.")
            return
        
        # Show current status
        col1, col2, col3 = st.columns(3)
        with col1:
            st.success(f"🤖 Model: {self.llm_manager.current_model}")
        with col2:
            st.success(f"📄 Documents: {len(self.documents)}")
        with col3:
            st.success("🔍 Vector search: Ready")
        
        # Initialize chat history
        if "messages" not in st.session_state:
            st.session_state.messages = []
        
        # Display chat messages
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
        
        # Chat input
        if prompt := st.chat_input("Ask about your documents..."):
            # Add user message to chat history
            st.session_state.messages.append({"role": "user", "content": prompt})
            with st.chat_message("user"):
                st.markdown(prompt)
            
            # Generate response
            with st.chat_message("assistant"):
                with st.spinner("Thinking..."):
                    try:
                        # Get relevant context from vector store
                        relevant_docs = self.vector_store.similarity_search(prompt)
                        
                        # Generate response
                        response = self.llm_manager.generate_response(prompt, relevant_docs)
                        st.markdown(response)
                        
                        # Add assistant response to chat history
                        st.session_state.messages.append({"role": "assistant", "content": response})
                        
                    except Exception as e:
                        error_msg = f"Error generating response: {str(e)}"
                        st.error(error_msg)
                        st.session_state.messages.append({"role": "assistant", "content": error_msg})
        
        # Quick action buttons
        st.markdown("---")
        st.markdown("**Quick Actions:**")
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            if st.button("📋 Summarize All"):
                summary_prompt = "Please provide a comprehensive summary of all the uploaded documents."
                st.session_state.messages.append({"role": "user", "content": summary_prompt})
                st.rerun()
        
        with col2:
            if st.button("🔍 Key Points"):
                key_points_prompt = "What are the most important key points across all documents?"
                st.session_state.messages.append({"role": "user", "content": key_points_prompt})
                st.rerun()
        
        with col3:
            if st.button("❓ Common Questions"):
                questions_prompt = "What questions would someone typically ask about these documents?"
                st.session_state.messages.append({"role": "user", "content": questions_prompt})
                st.rerun()
        
        with col4:
            if st.button("🗑️ Clear Chat"):
                st.session_state.messages = []
                st.rerun()
    
    def document_summary_section(self):
        """Document summary section"""
        if self.documents:
            st.header("📋 Document Summary")
            
            selected_doc = st.selectbox(
                "Select document to summarize:",
                list(self.documents.keys()),
                key="summary_doc_select"
            )
            
            if st.button("Generate Summary"):
                if self.llm_manager.current_model:
                    with st.spinner("Generating summary..."):
                        doc_text = self.documents[selected_doc]
                        summary_prompt = f"Please provide a comprehensive summary of the following document:\n\n{doc_text[:4000]}..."
                        summary = self.llm_manager.generate_response(summary_prompt)
                        st.markdown("### Summary:")
                        st.markdown(summary)
                else:
                    st.warning("Please select a model first.")
    
    def run(self):
        """Run the main application"""
        st.title("🤖 Local LLM Document Review")
        st.markdown("Review documents and PDFs using open source LLMs running locally")
        
        # Sidebar
        self.sidebar_model_management()
        
        # Main content
        col1, col2 = st.columns([1, 1])
        
        with col1:
            self.document_upload_section()
            self.document_summary_section()
        
        with col2:
            self.chat_interface()

def main():
    """Main function"""
    app = DocumentReviewApp()
    app.run()

if __name__ == "__main__":
    main()
